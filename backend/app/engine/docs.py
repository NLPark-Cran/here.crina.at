"""文档处理：PDF/DOCX/图片文本提取 + md 导出 docx/pdf"""
from __future__ import annotations

import asyncio
import base64
import logging
from pathlib import Path

log = logging.getLogger("crina.docs")

EXTRACT_MAX_CHARS = 20000  # 提取文本上限（防上下文爆炸）
VISION_MODEL = "qwen3-vl-plus"


async def extract_pdf(path: Path) -> str:
    """mutool 提取 PDF 文本"""
    proc = await asyncio.create_subprocess_exec(
        "mutool", "draw", "-F", "txt", "-o", "-", str(path),
        stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.DEVNULL,
    )
    out, _ = await asyncio.wait_for(proc.communicate(), timeout=60)
    if proc.returncode != 0:
        raise RuntimeError("mutool 提取失败")
    return out.decode("utf-8", errors="replace")


async def extract_docx(path: Path) -> str:
    def _run() -> str:
        import docx
        d = docx.Document(str(path))
        return "\n".join(p.text for p in d.paragraphs if p.text.strip())
    return await asyncio.to_thread(_run)


async def extract_image(path: Path, api_key: str) -> str:
    """视觉模型转写/描述图片内容"""
    from . import tokendance
    b64 = base64.b64encode(await asyncio.to_thread(path.read_bytes)).decode()
    suffix = "jpeg" if path.suffix.lower() in (".jpg", ".jpeg") else "png"
    ctx = [{
        "role": "user",
        "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/{suffix};base64,{b64}"}},
            {"type": "text", "text": "请详细转写这张图里的所有文字，并描述画面内容（用于后续文档检索引用）。"
                                    "有文字先逐字转写，再用一两句话描述画面。"},
        ],
    }]
    return await tokendance.chat_once(ctx, model=VISION_MODEL, temperature=0.2, max_tokens=2000)


async def extract(path: Path, kind: str, api_key: str) -> str:
    if kind == "pdf":
        text = await extract_pdf(path)
    elif kind == "docx":
        text = await extract_docx(path)
    else:
        text = await extract_image(path, api_key)
    return text.strip()[:EXTRACT_MAX_CHARS]


# ---------- md → docx / pdf 导出 ----------

def _md_lines(text: str) -> list[tuple[int, str]]:
    """把 md 行解析为 (标题级别 0=正文, 文本)，去掉常见行内标记"""
    import re
    out = []
    for raw in text.splitlines():
        line = raw.rstrip()
        level = 0
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            line = m.group(2)
        line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        line = re.sub(r"^[-*]\s+", "• ", line)
        line = re.sub(r"^>\s?", "▏", line)
        out.append((level, line))
    return out


def make_docx(text: str, title: str) -> bytes:
    import io

    import docx
    d = docx.Document()
    if title:
        d.add_heading(title, level=0)
    for level, line in _md_lines(text):
        if not line.strip():
            continue
        if level > 0:
            d.add_heading(line, level=min(level, 4))
        else:
            d.add_paragraph(line)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def make_pdf(text: str, title: str) -> bytes:
    import io

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas
    # 内置 CID 中文字体（免嵌入，主流阅读器都能显示）
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    _, height = A4
    left, top, leading = 20 * mm, height - 20 * mm, 7 * mm
    y = top
    c.setFont("STSong-Light", 16)
    if title:
        c.drawString(left, y, title[:60])
        y -= leading * 1.6
    c.setFont("STSong-Light", 10.5)
    max_chars = 44  # A4 一行大约这么多全角字
    for level, line in _md_lines(text):
        if level > 0:
            line = ("#" * level) + " " + line
        if not line.strip():
            y -= leading * 0.6
            continue
        # 简单全角折行
        while line:
            seg, line = line[:max_chars], line[max_chars:]
            if y < 20 * mm:
                c.showPage()
                c.setFont("STSong-Light", 10.5)
                y = top
            c.drawString(left, y, seg)
            y -= leading
    c.save()
    return buf.getvalue()
